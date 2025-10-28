import streamlit as st
import anthropic
import os
import json
import time
from pathlib import Path
from anthropic.types.message_create_params import MessageCreateParamsNonStreaming
from anthropic.types.messages.batch_create_params import Request
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ============ CONFIGURATION ============
BATCH_STATE_FILE = "batch_state.json"
COST_REPORT_FILE = "cost_report.txt"
MAX_INPUT_TOKENS = 195000
MAX_TOKENS = 64000
# =======================================

# ============ PRICING (per million tokens) ============
PRICING = {
    'input_below_200k': 1.50,
    'input_above_200k': 3.00,
    'output_below_200k': 7.50,
    'output_above_200k': 11.25,
}
# ======================================================

COMBINED_PROMPT = """You have TWO tasks to complete for this story:

TASK 1 - REWRITE THE STORY:

Create a narrator-ready version of this story with **approximately the same length as the original.**

PROCESS:(**YOU MUST FOLLOW THESE STEPS IN SAME ORDER AS MENTIONED**)
0.***MUST TO WRITE UP TO 60,000 TO 70,000 CHARACTERS TO CREATE A COMPLETE, ENGAGING NARRATIVE***
1. You must read and understand the full story arc first
2. **STORY QUALITY CHECK:** (**MUST TO QUALITY CHECK THE STORY FIRST**):
   - If the story is vague, unclear, lacks proper ending, or feels incomplete:
     * Expand and enhance the story based on its core theme
     * Add necessary context, character development, and emotional depth
     * Create a proper, satisfying conclusion
     **WRITE UP TO 60,000 TO 70,000 CHARACTERS TO CREATE A COMPLETE, ENGAGING NARRATIVE**
     * Maintain the original theme and emotional core while improving clarity
   - If the story is already complete and clear:
     * Proceed with standard rewrite matching original length

2. Rewrite using different words while preserving:
   - Number of major scenes/events
   - General paragraph structure
   - Emotional pacing
3. Replace:
   - All character names → new Spanish names
   - All dialogue → rephrased naturally
   - All descriptions → similar length but different words
4. Style (**MUST TO FOLLOW STYLE**):
   - Third-person narration
   - **DO NOT USE HEADINGS, TITLES, OR BREAKS.**
   - **Paragraph format** 
   - storyteller voice

TARGET:(**MUST TO FOLLOW THIS TARGET:**)
**Match original length within 5% (shorter or longer is fine if story needs it)**

LANGUAGE: **SPANISH** ,


TASK 2 - CREATE YOUTUBE METADATA:
You are an expert YouTube content strategist. Based on this story, create:
1. A *viral YouTube title* (max 100 characters) — it should grab attention, evoke curiosity or emotion, and perfectly fit the story's main theme.
2. A *short thumbnail text* (max 400 characters) — it should be bold, emotional, dramatic, or thought-provoking. Make it visually catchy and aligned with the story's emotion or twist.
3. A short hook (max 3 to 4 words) — **it should emotionally hooks viewers and makes them curious to click and watch the video till the end. It should sound like a shocking moment or emotional twist from the story — similar to phrases like 'SE QUEDÓ EN SHOCK.', 'COMENZÓ A LLORAR.', or '¡SE QUEDÓ BOQUIABIERTO!'."**
4. A *2–3 line YouTube description* — briefly summarize the story in an emotional, reflective, or motivational tone. Encourage viewers to watch till the end.
5. A list of *10 relevant tags* (comma-separated) that fit the story's themes and genre.
Style guidelines for metadata:
- **Use emotional triggers (love, regret, betrayal, hope, karma, redemption)**
- **Make the title and thumbnail feel like they belong to a cinematic story**
- Avoid clickbait — keep it believable yet gripping
- **Include curiosity elements like twists, lessons, or moral surprises**
- **Keep all outputs in the same language as the story (Spanish)**

OUTPUT FORMAT:

Please structure your response EXACTLY like this:

===REWRITTEN_STORY===

[Your complete rewritten story here in paragraphs]

===METADATA===

TITLE: [your title]

THUMBNAIL: [your thumbnail text]

HOOK: [your hook text]

DESCRIPTION: [your description]

TAGS: [comma-separated list of relevant tags, max 10]

===END==="""



class StoryProcessor:
    def __init__(self, api_key):
        self.client = anthropic.Anthropic(api_key=api_key)
    
    def estimate_tokens(self, text):
        """Estimate token count"""
        return len(text) / 4
    
    def calculate_cost(self, input_tokens, output_tokens):
        """Calculate cost based on token usage"""
        input_cost = 0
        output_cost = 0
        
        if input_tokens <= 200000:
            input_cost = (input_tokens / 1_000_000) * PRICING['input_below_200k']
        else:
            input_cost = (input_tokens / 1_000_000) * PRICING['input_above_200k']
        
        if output_tokens <= 200000:
            output_cost = (output_tokens / 1_000_000) * PRICING['output_below_200k']
        else:
            output_cost = (output_tokens / 1_000_000) * PRICING['output_above_200k']
        
        return {
            'input_cost': input_cost,
            'output_cost': output_cost,
            'total_cost': input_cost + output_cost
        }
    
    def scan_transcripts_folder(self, project_path):
        """Scan project transcripts folders"""
        transcript_files = []
        project_path = Path(project_path)
        
        # Scan all channel folders
        for channel_folder in sorted(project_path.iterdir()):
            if not channel_folder.is_dir() or channel_folder.name in ['__pycache__', '.git']:
                continue
            
            transcripts_dir = channel_folder / "transcripts"
            if not transcripts_dir.exists():
                continue
            
            # Load metadata if exists
            metadata_file = transcripts_dir / "metadata.json"
            metadata = {}
            if metadata_file.exists():
                try:
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        metadata_list = json.load(f)
                        # Convert list to dict keyed by folder name
                        metadata = {item['folder']: item for item in metadata_list}
                except:
                    pass
            
            # Scan all numbered folders
            for story_folder in sorted(transcripts_dir.iterdir(), key=lambda x: int(x.name) if x.name.isdigit() else 0):
                if not story_folder.is_dir():
                    continue
                
                # Find transcript.txt file
                txt_file = story_folder / "transcript.txt"
                
                if txt_file.exists():
                    folder_num = story_folder.name
                    folder_meta = metadata.get(folder_num, {})
                    
                    # Check if already processed
                    already_processed = (story_folder / "story.txt").exists()
                    
                    transcript_files.append({
                        'path': str(txt_file),
                        'channel_name': channel_folder.name,
                        'folder_name': folder_num,
                        'folder_path': str(story_folder),
                        'video_title': folder_meta.get('title', f"Story {folder_num}"),
                        'views': folder_meta.get('views', 0),
                        'already_processed': already_processed
                    })
        
        return transcript_files
    
    def submit_batch(self, stories_data):
        """Submit batch request to Claude"""
        try:
            requests = []
            
            for idx, story in enumerate(stories_data):
                # Read transcript
                with open(story['path'], 'r', encoding='utf-8') as f:
                    transcript = f.read()
                
                custom_id = f"story_{idx}_combined"
                
                request = Request(
                    custom_id=custom_id,
                    params=MessageCreateParamsNonStreaming(
                        model="claude-sonnet-4-20250514",
                        max_tokens=MAX_TOKENS,
                        messages=[{
                            "role": "user",
                            "content": f"{COMBINED_PROMPT}\n\nHere is the story:\n\n{transcript}"
                        }]
                    )
                )
                requests.append(request)
            
            # Create batch
            message_batch = self.client.messages.batches.create(requests=requests)
            
            return message_batch, None
            
        except Exception as e:
            return None, str(e)
    
    def check_batch_status(self, batch_id):
        """Check batch processing status"""
        try:
            batch = self.client.messages.batches.retrieve(batch_id)
            return batch, None
        except Exception as e:
            return None, str(e)
    
    def retrieve_batch_results(self, batch_id):
        """Retrieve batch results with token tracking"""
        try:
            results = []
            token_data = {}
            
            for result in self.client.messages.batches.results(batch_id):
                results.append(result)
                
                # Extract token usage data
                if result.result.type == "succeeded":
                    custom_id = result.custom_id
                    usage = result.result.message.usage
                    
                    token_data[custom_id] = {
                        'input_tokens': usage.input_tokens,
                        'output_tokens': usage.output_tokens,
                        'cache_creation_input_tokens': getattr(usage, 'cache_creation_input_tokens', 0),
                        'cache_read_input_tokens': getattr(usage, 'cache_read_input_tokens', 0),
                        'service_tier': 'batch'
                    }
            
            return results, token_data, None
            
        except Exception as e:
            return None, None, str(e)
    
    def save_batch_state(self, batch_id, stories, status="processing", token_data=None, completion_time=None):
        """Save batch state to JSON file"""
        try:
            state = {
                'batch_id': batch_id,
                'submission_timestamp': datetime.now().isoformat(),
                'processing_status': status,
                'completion_timestamp': completion_time,
                'stories_metadata': [
                    {
                        'custom_id': f"story_{i}_combined",
                        'channel_name': s['channel_name'],
                        'folder_name': s['folder_name'],
                        'folder_path': s['folder_path'],
                        'video_title': s['video_title']
                    }
                    for i, s in enumerate(stories)
                ],
                'token_tracking': token_data or {},
                'total_costs': None
            }
            
            with open(BATCH_STATE_FILE, 'w', encoding='utf-8') as f:
                json.dump(state, f, indent=2, ensure_ascii=False)
            
            return True
        except Exception as e:
            print(f"Error saving batch state: {e}")
            return False
    
    def load_batch_state(self):
        """Load batch state from JSON file"""
        try:
            if os.path.exists(BATCH_STATE_FILE):
                with open(BATCH_STATE_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
            return None
        except Exception as e:
            print(f"Error loading batch state: {e}")
            return None
    
    def generate_cost_report(self):
        """Generate cost report from batch state JSON"""
        try:
            state = self.load_batch_state()
            if not state:
                return None, "No batch state found"
            
            token_tracking = state.get('token_tracking', {})
            if not token_tracking:
                return None, "No token data available"
            
            # Calculate totals
            total_input = sum(data['input_tokens'] for data in token_tracking.values())
            total_output = sum(data['output_tokens'] for data in token_tracking.values())
            total_cache_creation = sum(data.get('cache_creation_input_tokens', 0) for data in token_tracking.values())
            total_cache_read = sum(data.get('cache_read_input_tokens', 0) for data in token_tracking.values())
            
            # Calculate costs
            cost_info = self.calculate_cost(total_input, total_output)
            
            # Generate report
            report_lines = []
            report_lines.append("=" * 60)
            report_lines.append("CLAUDE BATCH API - COST REPORT")
            report_lines.append("=" * 60)
            report_lines.append(f"Batch ID: {state['batch_id']}")
            report_lines.append(f"Submission Time: {state['submission_timestamp']}")
            report_lines.append(f"Completion Time: {state.get('completion_timestamp', 'N/A')}")
            report_lines.append(f"Status: {state['processing_status']}")
            report_lines.append("-" * 60)
            report_lines.append("TOTAL STATISTICS:")
            report_lines.append(f"  - Total Requests: {len(token_tracking)}")
            report_lines.append(f"  - Total Input Tokens: {total_input:,}")
            report_lines.append(f"  - Total Output Tokens: {total_output:,}")
            report_lines.append(f"  - Cache Creation Tokens: {total_cache_creation:,}")
            report_lines.append(f"  - Cache Read Tokens: {total_cache_read:,}")
            report_lines.append("-" * 60)
            report_lines.append("PER-STORY BREAKDOWN:")
            report_lines.append("")
            
            # Per-story details
            for custom_id, data in token_tracking.items():
                story_cost = self.calculate_cost(data['input_tokens'], data['output_tokens'])
                report_lines.append(f"Story: {custom_id}")
                report_lines.append(f"  - Input Tokens: {data['input_tokens']:,}")
                report_lines.append(f"  - Output Tokens: {data['output_tokens']:,}")
                report_lines.append(f"  - Cache Creation: {data.get('cache_creation_input_tokens', 0):,}")
                report_lines.append(f"  - Cache Read: {data.get('cache_read_input_tokens', 0):,}")
                report_lines.append(f"  - Estimated Cost: ${story_cost['total_cost']:.4f}")
                report_lines.append("")
            
            report_lines.append("-" * 60)
            report_lines.append("COST BREAKDOWN:")
            report_lines.append(f"  - Input Cost: ${cost_info['input_cost']:.4f}")
            report_lines.append(f"  - Output Cost: ${cost_info['output_cost']:.4f}")
            report_lines.append(f"  - TOTAL COST: ${cost_info['total_cost']:.4f}")
            report_lines.append("=" * 60)
            
            report_text = "\n".join(report_lines)
            
            # Save report to file
            with open(COST_REPORT_FILE, 'w', encoding='utf-8') as f:
                f.write(report_text)
            
            return report_text, cost_info['total_cost']
            
        except Exception as e:
            return None, f"Error generating report: {str(e)}"
    
    def parse_response(self, response_text):
        """Parse Claude's response into story and metadata"""
        try:
            parts = response_text.split("===REWRITTEN_STORY===")
            if len(parts) < 2:
                return None, None
            
            remaining = parts[1].split("===METADATA===")
            if len(remaining) < 2:
                return None, None
            
            story = remaining[0].strip()
            
            metadata_part = remaining[1].split("===END===")[0].strip()
            
            metadata = {}
            current_key = None
            current_value = []
            
            for line in metadata_part.split('\n'):
                line = line.strip()
                if not line:
                    continue
                
                if line.startswith('TITLE:'):
                    if current_key and current_value:
                        metadata[current_key] = ' '.join(current_value).strip()
                    current_key = 'title'
                    current_value = [line.replace('TITLE:', '').strip()]
                elif line.startswith('THUMBNAIL:'):
                    if current_key and current_value:
                        metadata[current_key] = ' '.join(current_value).strip()
                    current_key = 'thumbnail'
                    current_value = [line.replace('THUMBNAIL:', '').strip()]
                elif line.startswith('HOOK:'):
                    if current_key and current_value:
                        metadata[current_key] = ' '.join(current_value).strip()
                    current_key = 'hook'
                    current_value = [line.replace('HOOK:', '').strip()]
                elif line.startswith('DESCRIPTION:'):
                    if current_key and current_value:
                        metadata[current_key] = ' '.join(current_value).strip()
                    current_key = 'description'
                    current_value = [line.replace('DESCRIPTION:', '').strip()]
                elif line.startswith('TAGS:'):
                    if current_key and current_value:
                        metadata[current_key] = ' '.join(current_value).strip()
                    current_key = 'tags'
                    current_value = [line.replace('TAGS:', '').strip()]
                else:
                    if current_key:
                        current_value.append(line)
            
            if current_key and current_value:
                metadata[current_key] = ' '.join(current_value).strip()
            
            return story, metadata
            
        except Exception as e:
            print(f"Parse error: {e}")
            return None, None
    
    def create_word_document(self, story_text, metadata, output_path):
        """Create formatted Word document"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(metadata.get('title', 'Untitled Story'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add hook
            if metadata.get('hook'):
                hook_para = doc.add_paragraph()
                hook_run = hook_para.add_run(f"🎬 {metadata['hook']}")
                hook_run.font.size = Pt(14)
                hook_run.font.bold = True
                hook_run.font.color.rgb = RGBColor(220, 20, 60)
                hook_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            
            # Add story
            story_paragraphs = story_text.split('\n\n')
            for para_text in story_paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para_format = para.paragraph_format
                    para_format.line_spacing = 1.5
                    para_format.space_after = Pt(12)
                    
                    for run in para.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Calibri'
            
            doc.add_page_break()
            
            # Add metadata section
            doc.add_heading('📊 Metadata para YouTube', 1)
            
            sections = [
                ('📌 Título', metadata.get('title', '')),
                ('🖼️ Texto del Thumbnail', metadata.get('thumbnail', '')),
                ('🎯 Hook', metadata.get('hook', '')),
                ('📝 Descripción', metadata.get('description', '')),
                ('🏷️ Tags', metadata.get('tags', ''))
            ]
            
            for section_title, content in sections:
                if content:
                    heading = doc.add_heading(section_title, 2)
                    heading.runs[0].font.color.rgb = RGBColor(0, 102, 204)
                    
                    para = doc.add_paragraph(content)
                    para.runs[0].font.size = Pt(11)
                    doc.add_paragraph()
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error creating document: {e}")
            return False
    
    def save_results(self, stories_data, results, token_data):
        """Save processing results to files"""
        saved_count = 0
        total_cost = 0
        
        for result in results:
            if result.result.type == "succeeded":
                custom_id = result.custom_id
                
                # Extract index from custom_id (format: story_{idx}_combined)
                try:
                    idx = int(custom_id.split('_')[1])
                    story_info = stories_data[idx]
                except (IndexError, ValueError):
                    continue
                
                response_text = result.result.message.content[0].text
                story_text, metadata = self.parse_response(response_text)
                
                if story_text and metadata:
                    # Original folder_path is like Projects/ChannelOne/transcripts/1
                    # We need to save to Projects/ChannelOne/Rewritten/1
                    original_story_folder = Path(story_info['folder_path'])
                    channel_path = original_story_folder.parent.parent  # Go up from '1' to 'transcripts', then to 'ChannelOne'
                    
                    rewritten_story_folder = channel_path / "Rewritten" / story_info['folder_name']
                    rewritten_story_folder.mkdir(parents=True, exist_ok=True)

                    # Save story.txt
                    story_file = rewritten_story_folder / f"Story_{story_info['folder_name']}.txt"
                    with open(story_file, 'w', encoding='utf-8') as f:
                        f.write(story_text)

                    # Save metadata.json
                    metadata_file = rewritten_story_folder / "metadata.json"
                    with open(metadata_file, 'w', encoding='utf-8') as f:
                        json.dump(metadata, f, indent=2, ensure_ascii=False)

                    # Create Word document
                    docx_file = rewritten_story_folder / f"Story_{story_info['folder_name']}.docx"
                    self.create_word_document(story_text, metadata, str(docx_file))

                    saved_count += 1
                    
                    # Calculate cost for this story
                    if custom_id in token_data:
                        tokens = token_data[custom_id]
                        cost = self.calculate_cost(tokens['input_tokens'], tokens['output_tokens'])
                        total_cost += cost['total_cost']
        
        return saved_count, total_cost


class StoryProcessorApp:
    def __init__(self):
        self.init_session_state()
    
    def init_session_state(self):
        """Initialize session state variables"""
        if 'sp_scanned_files' not in st.session_state:
            st.session_state.sp_scanned_files = []
        if 'sp_selected_files' not in st.session_state:
            st.session_state.sp_selected_files = []
        if 'sp_processing' not in st.session_state:
            st.session_state.sp_processing = False
        if 'sp_batch_id' not in st.session_state:
            st.session_state.sp_batch_id = None
        if 'sp_batch_stories' not in st.session_state:
            st.session_state.sp_batch_stories = []
        if 'sp_token_info' not in st.session_state:
            st.session_state.sp_token_info = {}
        if 'sp_completed' not in st.session_state:
            st.session_state.sp_completed = False
        if 'sp_cost_report' not in st.session_state:
            st.session_state.sp_cost_report = None
        if 'sp_total_cost' not in st.session_state:
            st.session_state.sp_total_cost = 0
    
    def submit_stories_to_claude(self, selected_stories):
        """Submit selected stories to Claude Batch API"""
        try:
            api_key = st.session_state.get('claude_api_key', '')
            if not api_key:
                st.error("❌ Claude API Key not configured!")
                return False
            
            processor = StoryProcessor(api_key)
            
            with st.spinner("Submitting batch to Claude..."):
                batch, error = processor.submit_batch(selected_stories)
                
                if error:
                    st.error(f"❌ Error submitting batch: {error}")
                    return False
                
                if batch:
                    # Save batch state immediately after submission
                    processor.save_batch_state(
                        batch_id=batch.id,
                        stories=selected_stories,
                        status="processing"
                    )
                    
                    st.session_state.sp_batch_id = batch.id
                    st.session_state.sp_batch_stories = selected_stories
                    st.session_state.sp_processing = True
                    st.session_state.sp_completed = False
                    
                    st.success(f"✅ Batch submitted successfully!")
                    st.info(f"📝 Batch ID: {batch.id}")
                    st.info(f"📊 Processing {len(selected_stories)} stories")
                    
                    return True
            
            return False
            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            return False
    
    def render(self):
        """Render the batch processor UI"""
        st.markdown("## 🤖 Claude Batch Story Processor")
        st.markdown("Process multiple stories efficiently using Claude's Batch API")
        st.markdown("---")
        
        # Scan and select section
        if not st.session_state.sp_processing and not st.session_state.sp_completed:
            st.markdown("### 🔍 Scan & Select Stories")
            
            col1, col2 = st.columns([3, 1])
            
            with col1:
                if st.button("🔍 Scan Transcripts Folder", width='stretch', type="primary", key="sp_scan_button"):
                    api_key = st.session_state.get('claude_api_key', '')
                    if not api_key:
                        st.error("❌ Claude API Key not configured!")
                    else:
                        with st.spinner("Scanning transcripts folder..."):
                            processor = StoryProcessor(api_key)
                            scanned = processor.scan_transcripts_folder(st.session_state.current_project_path)
                            st.session_state.sp_scanned_files = scanned
                            
                            if scanned:
                                st.success(f"✅ Found {len(scanned)} transcript files")
                                # Auto-select all
                                st.session_state.sp_selected_files = list(range(len(scanned)))
                                time.sleep(1)
                                st.rerun()
                            else:
                                st.warning("⚠️ No transcript files found")
            
            with col2:
                if st.session_state.sp_scanned_files:
                    if st.button("🔄 Re-scan", width='stretch', key="sp_rescan_button"):
                        api_key = st.session_state.get('claude_api_key', '')
                        with st.spinner("Re-scanning..."):
                            processor = StoryProcessor(api_key)
                            scanned = processor.scan_transcripts_folder(st.session_state.current_project_path)
                            st.session_state.sp_scanned_files = scanned
                            st.session_state.sp_selected_files = list(range(len(scanned)))
                            time.sleep(1)
                            st.rerun()
        
        # Show scanned files
        if st.session_state.sp_scanned_files and not st.session_state.sp_processing and not st.session_state.sp_completed:
            st.markdown("---")
            st.markdown("### 📋 Select Stories to Process")
            
            # Select/Deselect All buttons
            col1, col2, col3 = st.columns([1, 1, 2])
            
            with col1:
                if st.button("☑️ Select All", width='stretch', key="sp_select_all_button"):
                    st.session_state.sp_selected_files = list(range(len(st.session_state.sp_scanned_files)))
                    st.rerun()
            
            with col2:
                if st.button("☐ Deselect All", width='stretch', key="sp_deselect_all_button"):
                    st.session_state.sp_selected_files = []
                    st.rerun()
            
            with col3:
                st.info(f"Selected: **{len(st.session_state.sp_selected_files)}** / {len(st.session_state.sp_scanned_files)} stories")
            
            st.markdown("---")
            
            # Group by channel
            channels = {}
            for idx, story_info in enumerate(st.session_state.sp_scanned_files):
                channel_name = story_info['channel_name']
                if channel_name not in channels:
                    channels[channel_name] = []
                channels[channel_name].append((idx, story_info))
            
            # Display by channel
            for channel_name, stories in sorted(channels.items()):
                channel_label = f"📁 {channel_name} ({len(stories)} transcripts)"
                
                with st.expander(channel_label, expanded=True):
                    for idx, story_info in stories:
                        col1, col2, col3 = st.columns([0.5, 3, 1.5])
                        
                        with col1:
                            is_selected = idx in st.session_state.sp_selected_files
                            if st.checkbox("Select", value=is_selected, key=f"sp_select_{idx}", label_visibility="collapsed"):
                                if idx not in st.session_state.sp_selected_files:
                                    st.session_state.sp_selected_files.append(idx)
                            else:
                                if idx in st.session_state.sp_selected_files:
                                    st.session_state.sp_selected_files.remove(idx)
                        
                        with col2:
                            status = "✅" if story_info['already_processed'] else "⏳"
                            st.write(f"{status} **Folder {story_info['folder_name']}**: {story_info['video_title'][:60]}...")
                        
                        with col3:
                            if story_info.get('views'):
                                st.caption(f"👁️ {story_info['views']:,} views")
            
            st.markdown("---")
            
            # Process button
            selected_count = len(st.session_state.sp_selected_files)
            if selected_count > 0:
                if st.button(f"🚀 Process {selected_count} Stories with Claude Batch API", type="primary", width='stretch', key="sp_process_button"):
                    # Get selected stories
                    selected_stories = [
                        st.session_state.sp_scanned_files[i]
                        for i in st.session_state.sp_selected_files
                    ]
                    
                    if self.submit_stories_to_claude(selected_stories):
                        time.sleep(2)
                        st.rerun()
            else:
                st.warning("⚠️ Please select at least one story")
        
        # Processing status
        if st.session_state.sp_processing and st.session_state.sp_batch_id:
            st.markdown("---")
            st.markdown("### ⏳ Batch Processing Status")
            
            st.info(f"📝 Batch ID: {st.session_state.sp_batch_id}")
            
            # Manual check button
            if st.button("🔄 Check Status Now", width='stretch', key="sp_check_status"):
                api_key = st.session_state.get('claude_api_key', '')
                processor = StoryProcessor(api_key)
                
                with st.spinner("Checking batch status..."):
                    batch, error = processor.check_batch_status(st.session_state.sp_batch_id)
                    
                    if error:
                        st.error(f"❌ Error checking status: {error}")
                    elif batch:
                        st.write(f"**Status:** {batch.processing_status}")
                        st.write(f"**Requests:** {batch.request_counts.processing} processing, {batch.request_counts.succeeded} succeeded, {batch.request_counts.errored} errored")
                        
                        if batch.processing_status == "ended":
                            st.success("✅ Batch completed! Retrieving results...")
                            
                            results, token_data, error = processor.retrieve_batch_results(st.session_state.sp_batch_id)
                            
                            if error:
                                st.error(f"❌ Error retrieving results: {error}")
                            else:
                                # Update batch state with completion data
                                processor.save_batch_state(
                                    batch_id=st.session_state.sp_batch_id,
                                    stories=st.session_state.sp_batch_stories,
                                    status="completed",
                                    token_data=token_data,
                                    completion_time=datetime.now().isoformat()
                                )
                                
                                # Save results to files
                                saved_count, total_cost = processor.save_results(
                                    st.session_state.sp_batch_stories,
                                    results,
                                    token_data
                                )
                                
                                # Generate cost report
                                cost_report, report_total_cost = processor.generate_cost_report()
                                
                                st.session_state.sp_cost_report = cost_report
                                st.session_state.sp_total_cost = report_total_cost if report_total_cost else total_cost
                                
                                st.balloons()
                                st.success(f"✅ Successfully processed {saved_count} stories!")
                                st.info(f"💰 Estimated cost: ${st.session_state.sp_total_cost:.4f}")
                                
                                st.session_state.sp_processing = False
                                st.session_state.sp_completed = True
                                # Don't auto-rerun - keep results visible
            
            st.info("💡 Click 'Check Status Now' to see if your batch is complete")
        
        # Completed - Show results and cost report
        if st.session_state.sp_completed:
            st.markdown("---")
            st.success("✅ Batch processing completed!")
            
            # Display cost report
            if st.session_state.sp_cost_report:
                st.markdown("### 💰 Cost Report")
                
                # Show summary
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("Total Stories", len(st.session_state.sp_batch_stories))
                
                with col2:
                    st.metric("Total Cost", f"${st.session_state.sp_total_cost:.4f}")
                
                with col3:
                    avg_cost = st.session_state.sp_total_cost / len(st.session_state.sp_batch_stories) if st.session_state.sp_batch_stories else 0
                    st.metric("Avg Cost/Story", f"${avg_cost:.4f}")
                
                # Show detailed report in expander
                with st.expander("📊 View Detailed Cost Report", expanded=False):
                    st.text(st.session_state.sp_cost_report)
                
                # Download buttons
                st.markdown("### 📥 Download Options")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    # Download as TXT
                    st.download_button(
                        label="📄 Download Cost Report (TXT)",
                        data=st.session_state.sp_cost_report,
                        file_name=f"cost_report_{st.session_state.sp_batch_id}.txt",
                        mime="text/plain",
                        width='stretch'
                    )
                
                with col2:
                    # Download batch state as JSON
                    processor = StoryProcessor(st.session_state.get('claude_api_key', ''))
                    batch_state = processor.load_batch_state()
                    if batch_state:
                        st.download_button(
                            label="📊 Download Batch Data (JSON)",
                            data=json.dumps(batch_state, indent=2, ensure_ascii=False),
                            file_name=f"batch_state_{st.session_state.sp_batch_id}.json",
                            mime="application/json",
                            width='stretch'
                        )
            
            st.markdown("---")
            
            # Reset button
            if st.button("🔄 Process More Stories", width='stretch', key="sp_reset"):
                st.session_state.sp_scanned_files = []
                st.session_state.sp_selected_files = []
                st.session_state.sp_processing = False
                st.session_state.sp_batch_id = None
                st.session_state.sp_completed = False
                st.session_state.sp_cost_report = None
                st.session_state.sp_total_cost = 0
                st.session_state.sp_batch_stories = []
                st.session_state.sp_token_info = {}
                st.rerun()
    
    def run(self):
        """Alias for render() to maintain compatibility"""
        self.render()


# Example usage (for testing)
if __name__ == "__main__":
    st.set_page_config(page_title="Story Batch Processor", layout="wide")
    
    # Mock session state for testing
    if 'claude_api_key' not in st.session_state:
        st.session_state.claude_api_key = st.text_input("Enter Claude API Key", type="password")
    
    if 'current_project_path' not in st.session_state:
        st.session_state.current_project_path = st.text_input("Enter Project Path", value=".")
    
    if st.session_state.claude_api_key:
        ui = StoryProcessorApp()
        ui.render()